"""Genera el documento EXPLICACION.docx a partir del contenido planeado."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Estilos globales
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)
    return h


def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")


def add_numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_table(headers, rows, header_color="1F365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], header_color)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# =========================================================================
# TÍTULO
# =========================================================================
title = doc.add_heading("El proyecto explicado para alguien que no sabe nada del tema", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)

add_quote(
    "Documento pensado para familiares, amigos, profes de otras áreas, o cualquiera que abra "
    "el repositorio del proyecto y diga \"¿qué es todo esto?\". No hace falta saber "
    "programación, matemáticas avanzadas, ni qué es un algoritmo. Vamos paso a paso."
)

p = doc.add_paragraph()
p.add_run("Autores del proyecto: ").bold = True
p.add_run("María Paula Saavedra Carrascal · Brayan Steven León")
p = doc.add_paragraph()
p.add_run("Materia: ").bold = True
p.add_run("Investigación de Operaciones — Universidad Autónoma de Bucaramanga (UNAB), 2026-I")

# =========================================================================
# 1
# =========================================================================
add_heading("1. La idea en una sola frase", 1)
add_para(
    "Teníamos que escoger el grupo de cosas más barato posible de una lista de 500 opciones, "
    "con la condición de que entre todas las cosas escogidas se cubrieran 500 requisitos distintos."
)
add_para(
    "Eso es todo. El reto está en que hay millones y millones de combinaciones posibles, "
    "y no podemos probarlas todas a mano."
)

# =========================================================================
# 2
# =========================================================================
add_heading("2. Una analogía cotidiana para entenderlo", 1)
add_para(
    "Imagínate que vas a hacer un botiquín de viaje y en la farmacia te ofrecen 500 cajas "
    "distintas de medicinas combinadas. Cada caja:"
)
add_bullets([
    "Tiene un precio (algunas son baratas, otras caras).",
    "Contiene algunos medicamentos (no todos). Por ejemplo, la caja #17 trae paracetamol, "
    "ibuprofeno y suero. La caja #203 trae solo aspirina y antialérgico.",
])
add_para(
    "Tú necesitas cubrir 500 dolencias distintas (dolor de cabeza, fiebre, alergia, dolor de "
    "estómago, etc., hasta 500). Lo único que importa es que cada dolencia esté cubierta por "
    "al menos una caja que hayas comprado."
)
add_para("La pregunta del problema es:", bold=True)
add_quote(
    "¿Cuál es la combinación de cajas que cubre las 500 dolencias gastando lo menos posible?"
)
add_para(
    "Eso, en lenguaje matemático elegante, se llama Set Covering Problem "
    "(Problema de Cubrimiento de Conjuntos)."
)
add_para("En nuestro caso:", bold=True)
add_bullets([
    "Las \"cajas\" son 500 subconjuntos posibles de imágenes médicas.",
    "Las \"dolencias\" son 500 imágenes específicas del dataset HAM10000 "
    "(un banco de fotos de lesiones de piel famoso en medicina).",
    "Los \"precios\" están en un archivo de Excel que nos dio el profe.",
])

# =========================================================================
# 3
# =========================================================================
add_heading("3. ¿Por qué no podemos probar todas las combinaciones?", 1)
add_para(
    "Porque las combinaciones posibles de 500 cajas son 2^500, un número con más de 150 ceros. "
    "Para que te hagas una idea:"
)
add_bullets([
    "El universo tiene unos 10^80 átomos.",
    "2^500 es muchísimo más grande que eso.",
])
add_para(
    "Si tu computador probara mil millones de combinaciones por segundo desde el Big Bang "
    "hasta hoy, no habría alcanzado a probar ni el 0.000...01%."
)
add_para(
    "Por eso este problema es famoso: pertenece a una categoría llamada NP-difícil, que son "
    "los problemas que en teoría tienen solución pero en la práctica nadie sabe cómo "
    "encontrarla en tiempo razonable, salvo para casos chiquitos."
)

# =========================================================================
# 4
# =========================================================================
add_heading("4. Las tres formas en que lo atacamos", 1)
add_para(
    "Como no podemos probar todo, los matemáticos llevan 70 años inventando trucos. "
    "Nosotros usamos tres de los más conocidos:"
)

add_heading("Forma A — El \"ansioso\" (Algoritmo Greedy)", 2)
add_para("Es la forma más simple y rápida.")
add_para(
    "La idea: en cada paso, agarra la caja que cubra más dolencias nuevas por cada peso que "
    "cueste. Es como ir al supermercado y siempre meter al carrito el producto con la mejor "
    "relación cantidad/precio en el momento."
)
add_bullets([
    "Ventajas: súper rápido (menos de 1 centésima de segundo).",
    "Desventajas: se enfoca en lo inmediato y a veces se \"casa\" con cajas que parecían "
    "buenas pero al final fuerzan a comprar otra caja cara. Es como hacer mercado con hambre.",
])
add_para("Quién lo inventó: Václav Chvátal en 1979. Es el algoritmo clásico de referencia.")

add_heading("Forma B — El \"perfeccionista\" (Programación Lineal Entera con Branch & Bound)", 2)
add_para("Esta forma sí encuentra la respuesta perfecta, pero puede demorarse mucho.")
add_para(
    "La idea: formular el problema como un sistema de ecuaciones gigantes, donde cada caja es "
    "una variable que vale 0 (no la compras) o 1 (la compras). Después se usa un programa "
    "especializado (nosotros usamos uno llamado CBC) que va explorando inteligentemente el "
    "árbol de posibilidades, descartando ramas enteras cuando puede demostrar que no van a "
    "mejorar la mejor solución encontrada hasta el momento."
)
add_bullets([
    "Ventajas: te garantiza la respuesta óptima (si te da tiempo de terminar).",
    "Desventajas: para 500 variables puede tardar 20 minutos o más, y si lo cortas antes "
    "te da una respuesta aproximada pero no la mejor.",
])

add_heading("Forma C — El \"imitador de la evolución\" (Algoritmo Genético)", 2)
add_para("Esta es la más curiosa. Está inspirada en la biología y la teoría de Darwin.")
add_para("La idea:")
add_numbered([
    "Generas 150 soluciones al azar, todas válidas (cada una compra un grupo distinto de "
    "cajas). A esto se le llama la \"población inicial\".",
    "Las pones a \"competir\": las que cuestan menos son las más \"fuertes\".",
    "Las soluciones fuertes se aparean entre sí — sí, literalmente: agarras dos buenas y "
    "combinas sus elecciones de cajas para crear una \"hija\". Esto se llama cruce.",
    "De vez en cuando una solución tiene una mutación al azar (cambia una caja por otra). "
    "Eso evita que toda la población se parezca demasiado.",
    "Repites por 500 generaciones. La población mejora poco a poco, como la evolución biológica.",
])
add_para(
    "Suena raro, pero funciona increíblemente bien en problemas donde lo perfecto es "
    "imposible y lo bueno alcanza."
)
add_para(
    "Quién lo inventó (la versión específica que usamos): Beasley y Chu en 1996, dos "
    "investigadores del Imperial College de Londres."
)

# =========================================================================
# 5
# =========================================================================
add_heading("5. Decisiones importantes que tomamos y por qué", 1)
add_para("A lo largo del proyecto hubo que decidir muchas cosas. Aquí las más relevantes:")

add_heading("Decisión 1 — ¿Por qué usar tres métodos y no uno solo?", 2)
add_para("Porque cada uno responde una pregunta distinta:")
add_bullets([
    "El greedy nos dice qué tan buena es la solución más obvia y rápida. Sirve como punto de "
    "comparación: si los otros métodos no la superan, no valen la pena.",
    "El ILP nos da la verdad absoluta. Sin él no sabríamos si nuestras soluciones aproximadas "
    "están cerca o lejos del óptimo.",
    "El genético nos da una solución casi óptima en mucho menos tiempo. En problemas reales "
    "(con 10,000 o 100,000 variables) el ILP ya no es viable y el genético sería la única "
    "opción práctica.",
])
add_para("Compararlos los tres es lo que le da rigor al trabajo.")

add_heading("Decisión 2 — ¿Qué tamaño de población usar en el genético?", 2)
add_para(
    "Probamos varios. Quedamos con 150 individuos por generación. Más grande = más calidad "
    "pero más lento. 150 fue el punto donde la calidad ya no mejoraba significativamente pero "
    "el tiempo todavía era razonable (~1 minuto)."
)

add_heading("Decisión 3 — ¿Cuántas generaciones?", 2)
add_para(
    "500 generaciones. Verificamos en las gráficas de convergencia que después de la "
    "generación 300-400 la mejora era prácticamente cero. 500 nos da un margen cómodo."
)

add_heading("Decisión 4 — La probabilidad de mutación: ¿alta o baja?", 2)
add_para(
    "Decidimos hacerla adaptativa: empieza alta (3%) y baja a lo largo del tiempo (0.5%). "
    "La intuición es que al principio nos interesa explorar mucho (mutar bastante) y al final "
    "ya queremos afinar (mutar poquito para no romper lo bueno encontrado)."
)

add_heading("Decisión 5 — El operador de reparación", 2)
add_para("Esta fue la decisión más importante de todas.", bold=True)
add_para(
    "Cuando dos soluciones se aparean, la \"hija\" muchas veces sale inválida: no cubre las "
    "500 dolencias, o compra cajas innecesarias. Si dejábamos esas soluciones inválidas, el "
    "algoritmo se quedaba estancado en resultados malos."
)
add_para("Implementamos un operador de reparación en dos fases:")
add_numbered([
    "Fase 1 — Hacerla válida: si falta cubrir alguna dolencia, agregar la caja más barata "
    "que la cubra.",
    "Fase 2 — Hacerla eficiente: si hay cajas redundantes (cuya remoción no destapa ninguna "
    "dolencia), quitarlas.",
])
add_para(
    "Antes de tener este reparador, el algoritmo nos daba resultados con error del 5% al 10% "
    "respecto al óptimo. Después de implementarlo, el error bajó por debajo del 2%. "
    "Hizo toda la diferencia."
)

add_heading("Decisión 6 — Probar con varias \"semillas\" al azar", 2)
add_para(
    "Los algoritmos genéticos usan números aleatorios, así que cada vez que los corres pueden "
    "dar un resultado un poquito distinto. Para asegurarnos de que no tuvimos suerte una sola "
    "vez, lo corrimos con 5 semillas diferentes (7, 13, 21, 42 y 99) y reportamos el promedio. "
    "Esto se llama análisis de robustez y es lo que se espera en un trabajo académico serio."
)

# =========================================================================
# 6
# =========================================================================
add_heading("6. Los resultados (en español plano)", 1)
add_table(
    ["Método", "Costo", "Cajas", "Tiempo"],
    [
        ["Cota inferior teórica (LP)", "$27,859.93", "—", "< 1 segundo"],
        ["Greedy de Chvátal", "$52,063", "23", "0.004 segundos"],
        ["Algoritmo Genético (semilla 13)", "$49,051", "22", "56 segundos"],
        ["ILP refinado (CBC + B&B)", "$49,051", "22", "20 minutos"],
    ],
)

add_heading("¿Qué nos dicen estos números?", 2)
add_numbered([
    "El greedy es rapidísimo pero se queda corto: gasta $3,000 más que la mejor solución y "
    "compra una caja extra.",
    "El genético logró la respuesta perfecta en menos de un minuto. Esto es asombroso "
    "considerando que el ILP exacto necesitó 20 minutos para confirmar que sí, ese era el "
    "óptimo verdadero.",
    "La cota teórica de $27,859 es un \"piso\" matemático: sabemos que ninguna solución "
    "entera (de comprar/no comprar cajas) puede bajar de ahí. Sirve para medir qué tan lejos "
    "estamos del límite absoluto.",
])

add_heading("Robustez del genético (5 corridas)", 2)
add_table(
    ["Semilla", "Costo", "Cajas"],
    [
        ["7", "$50,511", "23"],
        ["13", "$49,051", "22"],
        ["21", "$50,253", "23"],
        ["42", "$50,546", "23"],
        ["99", "$50,795", "23"],
        ["Promedio", "$50,231", "—"],
        ["Desviación", "$687 (1.37%)", "—"],
    ],
)
add_para(
    "La desviación del 1.37% significa que aun en la peor corrida, el genético da una "
    "respuesta bastante decente. Es un algoritmo confiable, no depende de tener suerte."
)

# =========================================================================
# 7
# =========================================================================
add_heading("7. El hallazgo más interesante del proyecto", 1)
add_para(
    "Esto fue una sorpresa y terminó siendo el aporte original del trabajo:"
)
add_para(
    "Cuando corrimos el ILP estándar (el \"perfeccionista\") con su configuración por defecto "
    "— 10 minutos de tiempo límite y una tolerancia de error del 0.01% —, nos dijo orgulloso: "
    "\"He encontrado la solución óptima: $50,123\" y nos dio una lista de 22 cajas."
)
add_para(
    "Pero después corrimos el genético y este encontró otra solución de $49,051 que también "
    "cubría todo y también con 22 cajas. Eso es $1,072 más barato que la supuesta \"solución "
    "óptima\" del ILP."
)
add_para("¿Cómo es posible que un algoritmo aproximado le gane a uno exacto?", bold=True)
add_para(
    "Respuesta: porque el ILP nunca llegó a terminar. Con 10 minutos no le alcanzó. Cuando "
    "dijo \"óptimo\" en realidad estaba diciendo \"óptimo dentro de mi tolerancia "
    "configurada\", y esa tolerancia era lo suficientemente laxa como para aceptar una "
    "respuesta sub-óptima."
)
add_para("Para confirmarlo, volvimos a correr el ILP pero con:")
add_bullets([
    "20 minutos de tiempo (en lugar de 10).",
    "Tolerancia del 0% absoluto (sin margen de error).",
    "Le pasamos el resultado del genético como pista inicial (warm-start).",
])
add_para(
    "Esta vez sí terminó completamente y confirmó: $49,051 es el óptimo entero verdadero."
)

add_heading("¿Por qué esto es importante?", 2)
add_para(
    "Porque demuestra una lección práctica que la gente suele olvidar:"
)
add_quote(
    "Un algoritmo \"exacto\" mal configurado puede dar peores respuestas que un algoritmo "
    "\"aproximado\" bien diseñado."
)
add_para(
    "En la vida real, donde los problemas son enormes y la gente confía ciegamente en "
    "herramientas como Excel Solver, Gurobi o CPLEX con sus configuraciones por defecto, este "
    "descuido puede costar miles o millones de pesos en decisiones reales (logística, "
    "transporte, asignación de recursos)."
)
p = doc.add_paragraph()
p.add_run("Conclusión académica del proyecto: ").bold = True
p.add_run(
    "validar siempre los resultados de un solver exacto contra una heurística, especialmente "
    "cuando el problema es denso y grande."
)

# =========================================================================
# 8
# =========================================================================
add_heading("8. ¿Para qué sirve esto en la vida real?", 1)
add_para(
    "El Set Covering Problem no es un juego matemático abstracto. Aparece todo el tiempo:"
)
add_bullets([
    "Salud pública: ¿en qué barrios pongo centros de vacunación para cubrir toda la ciudad "
    "gastando lo menos posible?",
    "Logística: ¿qué bodegas alquilo para que mis productos lleguen a todos los destinos?",
    "Aerolíneas: ¿qué tripulaciones asigno para cubrir todos los vuelos del mes con el "
    "mínimo costo? (Este es uno de los usos clásicos.)",
    "Telecomunicaciones: ¿dónde pongo antenas 5G para cubrir todo el territorio?",
    "Machine Learning: ¿qué subconjunto de imágenes uso para entrenar mi modelo de manera "
    "que vea todos los tipos de casos sin tener que procesarlas todas? "
    "(Justamente nuestro contexto con HAM10000.)",
])
add_para(
    "En todos estos casos, una mejora del 2% en costo puede traducirse en millones de pesos "
    "ahorrados al año."
)

# =========================================================================
# 9
# =========================================================================
add_heading("9. Cómo está organizado el código (por si alguien quiere mirar)", 1)
code_para = doc.add_paragraph()
code_run = code_para.add_run(
    "setcover-pau-brayan/\n"
    "├── data/                  # La instancia que nos dio el profe (no se toca)\n"
    "├── src/                   # El código en Python\n"
    "│   ├── loader.py          # Lee los archivos de entrada\n"
    "│   ├── greedy.py          # El método \"ansioso\"\n"
    "│   ├── exact.py           # El método \"perfeccionista\" (ILP)\n"
    "│   ├── ga.py              # El método \"evolutivo\" (Genético)\n"
    "│   ├── robustness.py      # Corrida con 5 semillas\n"
    "│   ├── compare.py         # Genera la tabla final\n"
    "│   └── figures.py         # Genera las gráficas del paper\n"
    "├── results/               # Los resultados en formato JSON\n"
    "├── paper/                 # El artículo académico en LaTeX (estilo IEEE)\n"
    "└── docs/                  # Este documento que estás leyendo"
)
code_run.font.name = "Consolas"
code_run.font.size = Pt(9)

add_heading("¿Cómo se reproduce todo?", 2)
add_para("Si alguien tiene Python instalado, puede correr todo el proyecto de una con:")
cmd_para = doc.add_paragraph()
cmd_run = cmd_para.add_run("bash run_all.sh")
cmd_run.font.name = "Consolas"
cmd_run.font.size = Pt(10)
add_para(
    "Y en aproximadamente 25 minutos tendrá todos los resultados, las gráficas y las tablas "
    "listas."
)

# =========================================================================
# 10
# =========================================================================
add_heading("10. Glosario rápido (por si te sonó a chino algo)", 1)
add_table(
    ["Término", "Qué significa en cristiano"],
    [
        ["NP-difícil",
         "Problema que en teoría se puede resolver pero en la práctica nadie sabe cómo "
         "hacerlo rápido."],
        ["Heurística",
         "Receta práctica que da buenas soluciones aunque no garantice la mejor."],
        ["Metaheurística",
         "Una heurística más sofisticada e inteligente (como el genético)."],
        ["Algoritmo exacto",
         "Receta que sí garantiza la mejor solución, si le das tiempo suficiente."],
        ["Cota inferior (LP)",
         "El piso teórico: ninguna solución real puede ser mejor que esto. Sirve como "
         "referencia."],
        ["Gap (brecha)",
         "Qué tan lejos está mi solución del óptimo. Un gap del 0% significa que encontré la "
         "mejor."],
        ["Semilla (seed)",
         "Un número que controla la \"aleatoriedad\" para que los experimentos sean "
         "reproducibles."],
        ["ILP", "Integer Linear Programming. Programación Lineal Entera."],
        ["GA", "Genetic Algorithm. Algoritmo Genético."],
        ["HAM10000",
         "Base de datos pública con 10,000 fotos de lesiones de piel, usada en investigación "
         "médica."],
        ["PuLP / CBC",
         "Herramientas de software libre que resuelven problemas de optimización."],
        ["Warm-start",
         "Darle una pista inicial al algoritmo para que no empiece desde cero."],
    ],
)

# =========================================================================
# 11
# =========================================================================
add_heading("11. ¿Qué falta?", 1)
add_para("A la fecha de este documento, falta:")
add_bullets([
    "Compilar el paper.tex en Overleaf y descargar el PDF final para entregar.",
    "Una app interactiva en Streamlit que permita a alguien probar el algoritmo en vivo "
    "(la rúbrica del profe pide \"interfaz visual\").",
    "Poner el nombre del profe en los reconocimientos del paper.",
])

# =========================================================================
# 12
# =========================================================================
add_heading("12. Cierre", 1)
add_para(
    "Este proyecto cubre un problema de optimización clásico (Set Covering) atacado con tres "
    "enfoques complementarios (greedy, exacto, genético), comparados rigurosamente y con un "
    "hallazgo metodológico relevante: la importancia de configurar bien los solvers exactos y "
    "validar siempre contra heurísticas."
)
add_para(
    "Si llegaste hasta acá y entendiste, ya sabes más de optimización combinatoria que el 99% "
    "de la gente. ¡Felicitaciones!"
)

add_quote(
    "Documento redactado en mayo de 2026 para acompañar el proyecto final de Investigación "
    "de Operaciones, UNAB."
)

out_path = "/home/nicolas/Documentos/setcover-pau-brayan/docs/EXPLICACION.docx"
doc.save(out_path)
print(f"OK -> {out_path}")
